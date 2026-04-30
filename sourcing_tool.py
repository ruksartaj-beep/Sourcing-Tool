import streamlit as st
import anthropic
import json
import io
import re
from datetime import datetime
from pathlib import Path
import pdfplumber
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.utils import get_column_letter

# ── Provider → Streamlit-secret key name mapping ──────────────────────────────
PROVIDER_SECRET_KEYS = {
    "Claude (Anthropic)": "ANTHROPIC_API_KEY",
    "ChatGPT (OpenAI)":   "OPENAI_API_KEY",
    "Groq (Free / Fast)": "GROQ_API_KEY",
    "Gemini (Google)":    "GEMINI_API_KEY",
    "Mistral AI":         "MISTRAL_API_KEY",
}

def _get_secret_key(provider: str) -> str:
    """Return the API key from st.secrets if the admin has set it, else empty string."""
    secret_name = PROVIDER_SECRET_KEYS.get(provider, "")
    try:
        return st.secrets.get(secret_name, "") or ""
    except Exception:
        return ""

# ── Provider catalogue ─────────────────────────────────────────────────────────
PROVIDERS = {
    "Claude (Anthropic)": {
        "icon": "🟠", "placeholder": "sk-ant-...", "help_url": "console.anthropic.com",
        "models": ["claude-sonnet-4-6", "claude-opus-4-6", "claude-haiku-4-5-20251001"],
        "model_help": "Sonnet = best balance · Opus = most thorough · Haiku = fastest",
    },
    "ChatGPT (OpenAI)": {
        "icon": "🟢", "placeholder": "sk-...", "help_url": "platform.openai.com",
        "models": ["gpt-4o", "gpt-4o-mini", "gpt-4-turbo", "gpt-3.5-turbo"],
        "model_help": "GPT-4o = best · GPT-4o-mini = cheapest smart · GPT-3.5 = budget",
    },
    "Groq (Free / Fast)": {
        "icon": "⚡", "placeholder": "gsk_...", "help_url": "console.groq.com",
        "models": ["llama-3.3-70b-versatile", "llama-3.1-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768"],
        "model_help": "Llama 3.3 70B = best · Llama 8B = fastest · Mixtral = alternative",
    },
    "Gemini (Google)": {
        "icon": "🔵", "placeholder": "AIza...", "help_url": "aistudio.google.com",
        "models": ["gemini-1.5-pro", "gemini-1.5-flash", "gemini-2.0-flash"],
        "model_help": "1.5 Pro = most capable · 1.5 Flash = fast · 2.0 Flash = latest",
    },
    "Mistral AI": {
        "icon": "🌀", "placeholder": "...", "help_url": "console.mistral.ai",
        "models": ["mistral-large-latest", "mistral-medium-latest", "mistral-small-latest", "open-mixtral-8x22b"],
        "model_help": "Large = most capable · Small = cheapest",
    },
}

# ── Page config ────────────────────────────────────────────────────────────────
st.set_page_config(page_title="AI Profile Ranker", page_icon="🎯", layout="wide",
                   initial_sidebar_state="expanded")

# ── Styling ────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  .block-container { padding-top: 1rem; }
  div[data-testid="stMetric"] {
    background: #f8fafc; border: 1px solid #e2e8f0;
    border-radius: 12px; padding: 14px; text-align: center;
  }
  div.stButton > button {
    background: linear-gradient(135deg, #2563eb, #7c3aed);
    color: white !important; border: none;
    border-radius: 10px; font-weight: 600; padding: 0.6rem 1.5rem;
  }
  div.stButton > button:hover { opacity: 0.9; }
  div.stDownloadButton > button {
    background: #059669; color: white !important;
    border: none; border-radius: 10px; font-weight: 600;
  }
  .skill-present  { color: #16a34a; font-weight: 600; }
  .skill-absent   { color: #dc2626; font-weight: 600; }
  .skill-partial  { color: #d97706; font-weight: 600; }
  .profile-card   { background: #f8fafc; border-radius: 10px; padding: 14px;
                    border: 1px solid #e2e8f0; margin-bottom: 10px; }
</style>
""", unsafe_allow_html=True)

# ── Config ─────────────────────────────────────────────────────────────────────
CONFIG_FILE = Path(__file__).parent / ".ranker_config.json"

def load_config():
    if CONFIG_FILE.exists():
        try:
            return json.loads(CONFIG_FILE.read_text())
        except Exception:
            return {}
    return {}

def save_config(data):
    try:
        CONFIG_FILE.write_text(json.dumps(data))
    except Exception:
        pass

# ── Text extraction ────────────────────────────────────────────────────────────
def extract_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            # Strategy 1: standard text extract
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""

            # Strategy 2: looser tolerances if too little text
            if len(page_text.strip()) < 30:
                page_text = page.extract_text(x_tolerance=6, y_tolerance=6) or ""

            # Strategy 3: word-by-word reconstruction if still empty
            if len(page_text.strip()) < 30:
                words = page.extract_words(x_tolerance=5, y_tolerance=5,
                                           keep_blank_chars=False, use_text_flow=True)
                page_text = " ".join(w["text"] for w in words)

            # Always extract tables too — many CVs store work history/education in tables
            try:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = "  |  ".join(
                            str(cell).strip() for cell in row
                            if cell and str(cell).strip()
                        )
                        if row_text and row_text not in page_text:
                            page_text += "\n" + row_text
            except Exception:
                pass

            text += page_text + "\n"

    return text.strip()

def extract_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract text from tables inside docx
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()

def extract_text(uploaded_file) -> str:
    file_bytes = uploaded_file.read()
    uploaded_file.seek(0)
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):  return extract_pdf(file_bytes)
    if name.endswith(".docx"): return extract_docx(file_bytes)
    return ""

def is_image_based_pdf(file_bytes: bytes) -> bool:
    """Returns True if PDF has almost no extractable text (likely scanned)."""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            total_chars = sum(
                len((page.extract_text() or "").strip())
                for page in pdf.pages[:3]
            )
        return total_chars < 100
    except Exception:
        return False

# ── LLM router ────────────────────────────────────────────────────────────────
def _call_llm(provider: str, api_key: str, model: str, prompt: str) -> str:
    if provider == "Claude (Anthropic)":
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(model=model, max_tokens=2400,
                                     messages=[{"role": "user", "content": prompt}])
        return msg.content[0].text.strip()
    elif provider == "ChatGPT (OpenAI)":
        from openai import OpenAI
        r = OpenAI(api_key=api_key).chat.completions.create(
            model=model, messages=[{"role": "user", "content": prompt}],
            max_tokens=2400, temperature=0.0)
        return r.choices[0].message.content.strip()
    elif provider == "Groq (Free / Fast)":
        from groq import Groq
        r = Groq(api_key=api_key).chat.completions.create(
            model=model, messages=[{"role": "user", "content": prompt}],
            max_tokens=2400, temperature=0.0)
        return r.choices[0].message.content.strip()
    elif provider == "Gemini (Google)":
        import google.generativeai as genai
        genai.configure(api_key=api_key)
        return genai.GenerativeModel(model).generate_content(prompt).text.strip()
    elif provider == "Mistral AI":
        from openai import OpenAI
        r = OpenAI(api_key=api_key, base_url="https://api.mistral.ai/v1").chat.completions.create(
            model=model, messages=[{"role": "user", "content": prompt}],
            max_tokens=2400, temperature=0.0)
        return r.choices[0].message.content.strip()
    else:
        raise ValueError(f"Unknown provider: {provider}")

# ── AI Ranking ────────────────────────────────────────────────────────────────
def rank_profile(api_key: str, provider: str, model: str,
                 jd_text: str, profile_text: str, file_name: str,
                 mandatory_reqs: str, additional_criteria: str,
                 skills_to_verify: list) -> dict:

    mandatory_section = ""
    if mandatory_reqs.strip():
        mandatory_section = f"""
=== MANDATORY REQUIREMENTS (NON-NEGOTIABLE) ===
The candidate MUST explicitly possess ALL of the following:
{mandatory_reqs}
RULES:
- If ANY is not clearly stated → mandatory_met: false
- If mandatory_met is false → verdict MUST be "Reject" and overall_score ≤ 25
- Absence = not present. List missing items in mandatory_missing.
"""

    additional_section = ""
    if additional_criteria.strip():
        additional_section = f"""
=== ADDITIONAL HIRING CRITERIA (carries 25% of overall score) ===
These are skills and expectations the hiring manager requires BEYOND what is written in the JD.
Evaluate the candidate against EACH item below and state clearly:
  - Meets: clearly evidenced in resume
  - Partially Meets: mentioned but limited depth/evidence
  - Does Not Meet: not found in resume
Score additional_criteria_match_pct as: (Meets × 1.0 + Partially Meets × 0.5) / total items × 100

{additional_criteria}
"""

    skills_section = ""
    if skills_to_verify:
        skill_lines = "\n".join(
            f"- {s['skill']}" + (f": {s['description']}" if s['description'] else "")
            for s in skills_to_verify
        )
        skills_section = f"""
=== SKILLS TO VERIFY ===
For EACH skill below, check the resume carefully and decide: Present / Partial / Absent.
"Present" only if clearly, explicitly stated with evidence. "Partial" if vaguely mentioned.
"Absent" if not mentioned at all. Include your finding in skills_checklist.
IMPORTANT: Search the ENTIRE resume including certifications tables, project descriptions,
technical skills sections, and summary. For certifications, check the certifications table
AND any badges or accreditation mentions anywhere in the document.

{skill_lines}
"""

    prompt = f"""You are a strict, senior Technical Recruiter and Hiring Manager.
Evaluate the candidate against the Job Description AND the Additional Hiring Criteria below.
The JD defines the baseline requirements. The Additional Hiring Criteria reflect what the
hiring manager specifically expects — treat both as equally important inputs to your evaluation.
Every gap and strength must be traceable back to the JD, the Additional Criteria, or the resume.
Do NOT invent requirements that appear in neither the JD nor the Additional Criteria.

{mandatory_section}
{additional_section}
{panel_section}
{skills_section}

════════════════════════════════════════════
PART A — CV PROFILE EXTRACTION
════════════════════════════════════════════
Extract factual information directly from the resume. Be thorough — read the ENTIRE document.

1. WORK HISTORY — list EVERY role in reverse chronological order:
   - company name, job title, start month/year, end month/year (or "Present")
   - Calculate duration_months precisely. If only years given (e.g. 2020–2021), estimate as 12 months.
   - short_tenure: set to TRUE if duration_months < 12. THIS IS MANDATORY — check every single role without exception.
   - Multiple roles at the same company count as one employer but list each role separately.
   IMPORTANT: Many CVs use a project-based or consulting format where experience appears as "Project #1", "Project #2", etc., or in table rows. Treat EACH project entry as a distinct role. Extract company, title, and dates from the table row or nearby text. Do NOT skip entries just because they are labelled "Project" instead of a traditional job title.

2. TOTAL EXPERIENCE — sum ALL duration_months across all roles (excluding gap periods), then express as years and months label.

3. JOB CHANGES — count of DISTINCT employers only (not role changes within the same company).

4. EDUCATION — search the ENTIRE resume carefully. Look for ANY of these keywords or section headings:
   Education · Educational Background · Academic Qualifications · Academics · Qualifications
   Degree · Bachelor · Master · MBA · PhD · B.Tech · B.E · B.Sc · M.Tech · M.Sc · MCA · BCA · Diploma · HSC · SSC
   University · College · Institute · School · Graduated · Year of Passing
   If found, extract: degree type, field of study, institution name, year.
   IMPORTANT: If you see any degree abbreviation anywhere in the resume (e.g. "B.Tech", "MBA", "M.Sc"), include it. Do NOT return an empty education list unless the resume truly has zero academic credentials anywhere.

5. CAREER GAPS — any gap > 3 months between consecutive roles:
   from date, to date, duration_months, any explanation stated in the resume.
   If no gaps exist, return an empty array [].

════════════════════════════════════════════
PART B — SKILL MATCH (JD-DRIVEN ONLY)
════════════════════════════════════════════
Read the Job Description carefully.

1. Identify must-have skills/requirements from the JD.
2. Identify good-to-have skills/requirements from the JD.
3. For each, check whether the candidate clearly has it, partially has it, or is missing it.
4. Only flag something as a gap if the JD actually requires or mentions it.
5. Only flag something as a strength if it is relevant to this JD.

════════════════════════════════════════════
PART C — SCORING PENALTIES (JD-BASED ONLY)
════════════════════════════════════════════
Start from 100. Apply penalties for gaps against the JD AND Additional Criteria:

- Each missing must-have skill from the JD: -10 pts each
- Experience below JD-required years by more than 2 years: -12 pts
- Role or domain clearly irrelevant to the JD: -15 pts
- Job hopping: if 2 or more DISTINCT employers had tenure < 12 months in the last 5 years: -10 pts
- Unexplained career gap > 6 months: -8 pts
- Each verified skill (from Skills to Verify) marked Absent that is also required by JD or Additional Criteria: -5 pts

DO NOT penalise for anything not in the JD or Additional Criteria.
Penalties stack. Be honest but fair.

════════════════════════════════════════════
PART D — OVERALL SCORING & VERDICT
════════════════════════════════════════════
Sub-scores (0–100 each):
- must_have_match_pct:           % of must-have JD requirements clearly present in resume
- good_to_have_match_pct:        % of good-to-have JD requirements present
- experience_match_pct:          relevance + depth + years vs JD requirement
- additional_criteria_match_pct: % match against Additional Hiring Criteria
                                 (Meets=100%, Partially Meets=50%, Does Not Meet=0% per item)
                                 Set to 100 only if no additional criteria were provided.

overall_score = (must_have × 0.35) + (experience × 0.25) + (good_to_have × 0.15) + (additional × 0.25)

Score guide:
  85–100: Exceptional — exceeds JD + additional criteria with strong evidence
  70–84:  Strong — meets nearly all requirements, minor gaps only
  55–69:  Average — meets core requirements but has clear gaps
  35–54:  Weak — significant gaps against JD or additional criteria
  0–34:   Poor — missing most requirements or failed mandatory check

VERDICT:
"Strong Select" → overall_score ≥ 80 AND must_have_match_pct ≥ 75 AND additional_criteria_match_pct ≥ 70
"Consider"      → overall_score ≥ 55 AND does not meet Strong Select threshold
"Reject"        → overall_score < 55 OR mandatory_met = false

CALIBRATION: Be honest — most candidates will not be a perfect fit. Expect 1–2 Strong Select per 10 resumes max.

CONFIDENCE: "High" = detailed resume, clear evidence. "Medium" = some inference needed. "Low" = sparse/ambiguous.

════════════════════════════════════════════
OUTPUT — Return ONLY valid JSON, no markdown:
════════════════════════════════════════════
{{
  "cv_profile": {{
    "total_experience_years": <number e.g. 7.5>,
    "total_experience_label": "<e.g. 7 years 6 months>",
    "job_changes_count": <integer — distinct employers>,
    "roles": [
      {{"company": "<name>", "title": "<title>", "start": "<Mon YYYY>", "end": "<Mon YYYY or Present>",
        "duration_months": <integer>, "short_tenure": <true if < 12 months>}}
    ],
    "education": [
      {{"degree": "<e.g. B.Tech>", "field": "<e.g. Computer Science>",
        "institution": "<name>", "year": "<e.g. 2016>"}}
    ],
    "career_gaps": [
      {{"from": "<Mon YYYY>", "to": "<Mon YYYY>", "duration_months": <integer>, "note": "<explanation or empty string>"}}
    ]
  }},
  "skills_checklist": [
    {{"skill": "<skill name>", "status": "<Present | Partial | Absent>",
      "evidence": "<brief quote or note from resume, or 'Not found'>"}}
  ],
  "overall_score": <integer 0-100>,
  "must_have_match_pct": <integer 0-100>,
  "good_to_have_match_pct": <integer 0-100>,
  "experience_match_pct": <integer 0-100>,
  "additional_criteria_match_pct": <integer 0-100>,
  "verdict": "<Strong Select | Consider | Reject>",
  "confidence_level": "<High | Medium | Low>",
  "seniority_alignment": "<Underqualified | Well-matched | Overqualified>",
  "role_relevance": "<Same | Similar | Different>",
  "domain_relevance": "<Aligned | Partially Aligned | Not Aligned>",
  "years_match": "<Exceeds | Meets | Partial | Below>",
  "penalties_applied": ["<penalty description with pts deducted>"],
  "key_strengths": ["<strength relevant to this JD>", "<strength 2>", "<strength 3>"],
  "key_gaps": ["<gap relevant to this JD>", "<gap 2>", "<gap 3>"],
  "risk_flags": ["<risk flag or 'None identified'>"],
  "additional_criteria_match": ["<evaluation per criterion>"],
  "final_explanation": "<2-3 sentences: score, verdict, key reasons — all tied to the JD>",
  "mandatory_met": <true | false>,
  "mandatory_missing": ["<missing mandatory item>"]
}}

════════════════════════════════════════════

JOB DESCRIPTION:
{jd_text[:5000]}

---

CANDIDATE PROFILE: {file_name}
{profile_text[:6000]}

Evaluate strictly against the JD above. Return only JSON."""

    raw = _call_llm(provider, api_key, model, prompt)
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", raw)
    if m:
        raw = m.group(1).strip()
    return json.loads(raw)


def _error_result(name: str, reason: str) -> dict:
    return {
        "file": name, "overall_score": 0,
        "must_have_match_pct": 0, "good_to_have_match_pct": 0,
        "experience_match_pct": 0, "additional_criteria_match_pct": 0,
        "verdict": "N/A", "confidence_level": "N/A",
        "seniority_alignment": "N/A", "role_relevance": "N/A",
        "domain_relevance": "N/A", "years_match": "N/A",
        "cv_profile": {}, "skills_checklist": [],
        "penalties_applied": [],
        "key_strengths": [], "key_gaps": [],
        "risk_flags": ["Analysis failed"],
        "additional_criteria_match": [],
        "final_explanation": f"Analysis failed: {reason}",
        "mandatory_met": False, "mandatory_missing": ["Analysis failed"],
    }


# ── Excel export ───────────────────────────────────────────────────────────────
def generate_excel(results: list, jd_name: str) -> io.BytesIO:
    wb = Workbook()
    ws = wb.active
    ws.title = "Rankings"

    headers = [
        "Rank", "Candidate",
        # CV Profile
        "Total Experience", "Job Changes", "Education", "Career Gaps",
        # Scores
        "Overall Score (%)", "Verdict", "Confidence", "Seniority",
        "Must-Have %", "Good-to-Have %", "Experience %", "Additional Criteria %",
        "Role Relevance", "Domain Relevance", "Years Match",
        # Assessment
        "Skills Checklist",
        "Penalties Applied", "Key Strengths", "Key Gaps", "Risk Flags",
        "Additional Criteria", "Final Explanation",
        "Mandatory Met", "Missing Mandatory",
    ]

    hdr_fill = PatternFill(start_color="2563EB", end_color="2563EB", fill_type="solid")
    hdr_font = Font(color="FFFFFF", bold=True, size=10)
    center   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    wrap_top = Alignment(wrap_text=True, vertical="top")

    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.fill = hdr_fill; c.font = hdr_font; c.alignment = center
    ws.row_dimensions[1].height = 28

    color_map = [(80,"D1FAE5"),(60,"FEF9C3"),(40,"FED7AA"),(0,"FEE2E2")]

    def row_color(score, mand_ok):
        if not mand_ok: return "FEE2E2"
        for t, c in color_map:
            if score >= t: return c
        return "FEE2E2"

    def yn(v):
        s = str(v).lower()
        if s == "yes": return "✅ Yes"
        if s == "no":  return "❌ No"
        if s == "partial": return "🟡 Partial"
        return v or "—"

    for i, r in enumerate(results):
        rn = i + 2
        score    = r.get("overall_score", 0)
        mand_ok  = r.get("mandatory_met", True)
        rf = PatternFill(start_color=row_color(score, mand_ok),
                         end_color=row_color(score, mand_ok), fill_type="solid")

        cname = r["file"]
        for ext in [".pdf",".docx",".PDF",".DOCX"]: cname = cname.replace(ext,"")

        cp = r.get("cv_profile", {})

        edu_str = "; ".join(
            f"{e.get('degree','')} {e.get('field','')} – {e.get('institution','')} ({e.get('year','')})"
            for e in cp.get("education", [])
        ) or "—"

        gaps_str = "; ".join(
            f"{g.get('from','')}–{g.get('to','')}: {g.get('duration_months','')} months"
            + (f" ({g['note']})" if g.get("note") else "")
            for g in cp.get("career_gaps", [])
        ) or "None"

        sc_str = "; ".join(
            f"{s['skill']}: {s['status']}" + (f" [{s.get('evidence','')}]" if s.get('evidence') else "")
            for s in r.get("skills_checklist", [])
        ) or "—"

        vals = [
            i+1, cname,
            cp.get("total_experience_label","—"), cp.get("job_changes_count","—"), edu_str, gaps_str,
            score, r.get("verdict",""), r.get("confidence_level",""), r.get("seniority_alignment",""),
            r.get("must_have_match_pct",""), r.get("good_to_have_match_pct",""),
            r.get("experience_match_pct",""), r.get("additional_criteria_match_pct",""),
            r.get("role_relevance",""), r.get("domain_relevance",""), r.get("years_match",""),
            sc_str,
            "; ".join(r.get("penalties_applied",[])),
            "; ".join(r.get("key_strengths",[])),
            "; ".join(r.get("key_gaps",[])),
            "; ".join(r.get("risk_flags",[])),
            "; ".join(r.get("additional_criteria_match",[])),
            r.get("final_explanation",""),
            "✅ Yes" if mand_ok else "❌ No",
            "; ".join(r.get("mandatory_missing",[])) or "—",
        ]

        for ci, v in enumerate(vals, 1):
            c = ws.cell(row=rn, column=ci, value=v)
            c.fill = rf; c.alignment = wrap_top
        ws.row_dimensions[rn].height = 90

    col_widths = [6, 24, 16, 12, 45, 35,
                  13, 14, 11, 16, 12, 14, 13, 18, 14, 17, 12,
                  55, 50, 55, 55, 40, 50, 65, 13, 38]
    for ci, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(ci)].width = w

    # Summary sheet
    ws2 = wb.create_sheet("Summary")
    total = len(results)
    avg   = round(sum(r.get("overall_score",0) for r in results)/total) if total else 0
    ss    = sum(1 for r in results if r.get("verdict")=="Strong Select")
    co    = sum(1 for r in results if r.get("verdict")=="Consider")
    re_   = sum(1 for r in results if r.get("verdict")=="Reject")
    fm    = sum(1 for r in results if not r.get("mandatory_met",True))

    tf = Font(bold=True, size=14, color="2563EB")
    bf = Font(bold=True)
    rows2 = [
        ("AI Profile Ranker — Evaluation Report",""),("",""),
        ("Generated On", datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Job Description", jd_name),("",""),
        ("Total Profiles Analyzed", total),("Average Overall Score", f"{avg}%"),
        ("Strong Select", ss),("Consider", co),("Reject", re_),
        ("Failed Mandatory Requirements", fm),
    ]
    for ri,(lab,val) in enumerate(rows2,1):
        ws2.cell(ri,1,lab); ws2.cell(ri,2,val)
        if ri==1: ws2.cell(ri,1).font=tf
        elif lab: ws2.cell(ri,1).font=bf
    ws2.column_dimensions["A"].width=34
    ws2.column_dimensions["B"].width=30

    out = io.BytesIO()
    wb.save(out); out.seek(0)
    return out


# ── Shared persistent store (survives browser refresh, shared across all users) ─
@st.cache_resource
def _store():
    return {"history": [], "jd_library": {}}


# ── Main App ──────────────────────────────────────────────────────────────────
def main():
    config = load_config()

    # ── Logo ──────────────────────────────────────────────────────────────────
    logo_path = Path(__file__).parent / "DN Logo.jpg"
    if logo_path.exists():
        st.logo(str(logo_path), size="large")

    # ── Header ────────────────────────────────────────────────────────────────
    if logo_path.exists():
        _, logo_col, _ = st.columns([1, 2, 1])
        with logo_col:
            st.image(str(logo_path), use_container_width=True)
    st.title("🎯 AI Profile Ranker")
    st.caption("Evaluate and rank candidates — powered by Claude, ChatGPT, Groq, Gemini, Mistral and more")
    st.divider()

    # ═══════════════════════════════════════════
    # SIDEBAR
    # ═══════════════════════════════════════════
    with st.sidebar:

        # ── AI Settings (collapsed by default) ──────────────────────────────
        with st.expander("⚙️ AI Settings", expanded=False):
            provider_names = list(PROVIDERS.keys())
            saved_provider = config.get("provider","ChatGPT (OpenAI)")
            if "provider" not in st.session_state:
                st.session_state.provider = saved_provider

            provider = st.selectbox(
                "AI Provider", options=provider_names,
                index=provider_names.index(st.session_state.provider)
                      if st.session_state.provider in provider_names else 0,
                format_func=lambda p: f"{PROVIDERS[p]['icon']}  {p}",
                key="provider_select",
            )
            st.session_state.provider = provider
            prov = PROVIDERS[provider]

            skey       = f"api_key_{provider}"
            secret_key = _get_secret_key(provider)

            if secret_key:
                # ── Key pre-configured by admin via Streamlit Secrets ──────────
                api_key = secret_key
                st.success("🔑 API key configured by admin — ready to use")
                st.caption("Your administrator has set this key. No action needed.")
            else:
                # ── Manual entry (local / self-hosted) ─────────────────────────
                saved_key = config.get(skey, "")
                if skey not in st.session_state:
                    st.session_state[skey] = saved_key

                api_key_input = st.text_input(
                    "API Key", value=st.session_state[skey],
                    type="password", placeholder=prov["placeholder"],
                    help=f"Get your key at {prov['help_url']}",
                )
                st.caption(f"🔑 [{prov['help_url']}](https://{prov['help_url']})")

                remember = st.checkbox("💾 Remember key on this computer", value=bool(saved_key))
                lkey = f"_last_{skey}"
                if api_key_input and api_key_input != st.session_state.get(lkey, ""):
                    st.session_state[skey] = api_key_input
                    st.session_state[lkey] = api_key_input
                    if remember:
                        cfg = load_config()
                        cfg[skey] = api_key_input
                        cfg["provider"] = provider
                        save_config(cfg)
                        st.success("✅ Saved")
                    else:
                        st.success("✅ Set for this session")
                elif api_key_input:
                    st.session_state[skey] = api_key_input

                api_key = st.session_state.get(skey, "")

            saved_model = config.get("model", prov["models"][0])
            model_index = prov["models"].index(saved_model) if saved_model in prov["models"] else 0
            model = st.selectbox("Model", options=prov["models"],
                                 index=model_index, help=prov["model_help"])

        st.divider()

        # ── Mandatory Requirements ───────────────────────────────────────────
        st.header("🔒 Mandatory Requirements")
        st.caption("Fill this **after** reviewing the JD. Missing ANY of these = automatic **Reject**.")
        mandatory_reqs = st.text_area(
            "One requirement per line — leave blank if none",
            placeholder=(
                "e.g. 5+ years of experience in the relevant role\n"
                "e.g. Specific certification required by JD\n"
                "e.g. Must-have tool or technology from JD"
            ),
            height=120,
        )
        if mandatory_reqs.strip():
            req_list = [r.strip() for r in mandatory_reqs.splitlines() if r.strip()]
            st.warning(f"⛔ {len(req_list)} mandatory req(s)")
            for r in req_list:
                st.markdown(f"  • {r}")
        else:
            st.caption("No mandatory requirements set — all candidates will be scored, none auto-rejected.")

        st.divider()

        # ── Skills to Verify ─────────────────────────────────────────────────
        st.header("🔍 Skills to Verify")
        st.caption("Add after reading the JD — AI marks each skill Present / Partial / Absent on every CV")

        if "skills_to_verify" not in st.session_state:
            st.session_state.skills_to_verify = []

        sv_col1, sv_col2 = st.columns([2, 3])
        with sv_col1:
            sv_skill = st.text_input("Skill name", placeholder="e.g. Relevant tool from JD",
                                     key="sv_skill_input")
        with sv_col2:
            sv_desc = st.text_input("What to look for (optional)",
                                    placeholder="e.g. hands-on project experience, not just listed",
                                    key="sv_desc_input")

        if st.button("➕ Add Skill", key="add_skill_btn"):
            if sv_skill.strip():
                st.session_state.skills_to_verify.append(
                    {"skill": sv_skill.strip(), "description": sv_desc.strip()}
                )
                st.rerun()
            else:
                st.warning("Enter a skill name first.")

        if st.session_state.skills_to_verify:
            for i, sv in enumerate(st.session_state.skills_to_verify):
                sa, sb = st.columns([4, 1])
                with sa:
                    label = sv["skill"]
                    if sv["description"]:
                        label += f" — *{sv['description'][:40]}*"
                    st.caption(f"• {label}")
                with sb:
                    if st.button("✕", key=f"del_skill_{i}"):
                        st.session_state.skills_to_verify.pop(i)
                        st.rerun()

        st.divider()

        # ── Additional Criteria ──────────────────────────────────────────────
        st.header("🎯 Additional Criteria")
        st.caption("Hiring manager expectations beyond the JD — carries **25% of the overall score**")
        additional_criteria = st.text_area(
            "Free text",
            placeholder=(
                "Prefer product companies\n"
                "Avoid frequent job hoppers\n"
                "Must have stakeholder management\n"
                "Prefer FinTech / BFSI background"
            ),
            height=130,
        )
        if additional_criteria.strip():
            st.info("📌 Will factor into scoring")

        st.divider()
        st.caption("🔐 Keys stored locally only. Never shared.")

    # ═══════════════════════════════════════════
    # SHARED STORE (persists across refreshes)
    # ═══════════════════════════════════════════
    shared = _store()

    # ═══════════════════════════════════════════
    # MAIN AREA
    # ═══════════════════════════════════════════
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📋 Step 1: Job Description")

        # ── Load from JD Library ──────────────────────────────────────────────
        jd_library = shared["jd_library"]
        if jd_library:
            lib_options = ["— New JD —"] + list(jd_library.keys())
            selected_jd = st.selectbox("📂 Load saved JD", lib_options, key="jd_library_select")
        else:
            selected_jd = "— New JD —"
            st.caption("💡 No saved JDs yet — save one below after entering it.")

        jd_mode = st.radio("Input method", ["Upload file (PDF/DOCX)", "Paste JD text"],
                           horizontal=True, key="jd_input_mode")
        jd_file = None
        jd_pasted_text = ""
        jd_pasted_name = "Pasted JD"

        # Pre-fill from library if selected
        prefill_text = jd_library.get(selected_jd, "") if selected_jd != "— New JD —" else ""
        prefill_name = selected_jd if selected_jd != "— New JD —" else "Pasted JD"

        if jd_mode == "Upload file (PDF/DOCX)":
            jd_file = st.file_uploader("Upload JD", type=["pdf","docx"], key="jd_uploader")
            if jd_file:
                st.success(f"✅ {jd_file.name}")
                # Allow saving uploaded JD to library
                save_jd_name = st.text_input("Save to JD Library as (optional)",
                    value=jd_file.name.rsplit(".",1)[0],
                    placeholder="e.g. Senior Data Engineer",
                    key="save_jd_name_input")
                if st.button("💾 Save this JD to Library", key="save_jd_btn"):
                    if save_jd_name.strip():
                        jd_bytes = jd_file.read(); jd_file.seek(0)
                        extracted = extract_pdf(jd_bytes) if jd_file.name.lower().endswith(".pdf") else extract_docx(jd_bytes)
                        if extracted.strip():
                            shared["jd_library"][save_jd_name.strip()] = extracted.strip()
                            st.success(f"✅ Saved '{save_jd_name.strip()}' to JD Library")
                            st.rerun()
                        else:
                            st.warning("Could not extract text from this file to save.")
        else:
            jd_pasted_name = st.text_input("JD title",
                value=prefill_name if prefill_name != "Pasted JD" else "",
                placeholder="e.g. Senior Data Engineer – Databricks", key="jd_paste_name") or "Pasted JD"
            jd_pasted_text = st.text_area("Paste Job Description",
                value=prefill_text,
                placeholder="Paste the full JD here…", height=200, key="jd_paste_text")

            if jd_pasted_text.strip():
                st.success(f"✅ JD ready ({len(jd_pasted_text):,} chars)")
                sa_col, sb_col = st.columns([3, 2])
                with sa_col:
                    save_jd_name = st.text_input("Save to JD Library as",
                        value=jd_pasted_name if jd_pasted_name != "Pasted JD" else "",
                        placeholder="e.g. Senior Data Engineer",
                        key="save_jd_name_input")
                with sb_col:
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("💾 Save to Library", key="save_jd_btn", use_container_width=True):
                        name_to_save = save_jd_name.strip() or jd_pasted_name
                        shared["jd_library"][name_to_save] = jd_pasted_text.strip()
                        st.success(f"✅ Saved '{name_to_save}'")
                        st.rerun()
                if selected_jd != "— New JD —":
                    if st.button(f"🗑️ Remove '{selected_jd}' from library", key="del_jd_btn"):
                        del shared["jd_library"][selected_jd]
                        st.rerun()

    with col2:
        st.subheader("👥 Step 2: Candidate Profiles")
        profile_files = st.file_uploader(
            "Upload profiles (select multiple)",
            type=["pdf","docx"], accept_multiple_files=True, key="profiles_uploader"
        )
        if profile_files:
            st.success(f"✅ {len(profile_files)} profile(s) ready")
            for pf in profile_files:
                st.caption(f"  • {pf.name}")

    # ── Paste resume ──────────────────────────────────────────────────────────
    if "pasted_profiles" not in st.session_state:
        st.session_state.pasted_profiles = []

    with st.expander("📝  Step 3 (Optional): Paste Resume Text"):
        st.caption("Add a candidate by pasting their resume text directly.")
        pc1, pc2 = st.columns([1,3])
        with pc1:
            paste_name = st.text_input("Candidate name", placeholder="e.g. Jane Smith",
                                       key="paste_name_input")
        with pc2:
            paste_text = st.text_area("Resume text", placeholder="Paste full resume here…",
                                      height=200, key="paste_text_input")
        if st.button("➕ Add Pasted Resume", key="add_paste_btn"):
            if paste_text.strip():
                n = paste_name.strip() or f"Pasted Resume {len(st.session_state.pasted_profiles)+1}"
                st.session_state.pasted_profiles.append({"name": n, "text": paste_text.strip()})
                st.rerun()
            else:
                st.warning("Paste some resume text first.")
        if st.session_state.pasted_profiles:
            st.markdown(f"**{len(st.session_state.pasted_profiles)} pasted resume(s) queued:**")
            for i, pr in enumerate(st.session_state.pasted_profiles):
                ra, rb = st.columns([5,1])
                with ra: st.caption(f"• {pr['name']}  ({len(pr['text']):,} chars)")
                with rb:
                    if st.button("🗑️", key=f"rm_paste_{i}"):
                        st.session_state.pasted_profiles.pop(i); st.rerun()


    st.divider()

    # ── Resolve effective JD from library if nothing else provided ────────────
    if selected_jd != "— New JD —" and not jd_file and not jd_pasted_text.strip():
        jd_pasted_text = shared["jd_library"].get(selected_jd, "")
        jd_pasted_name = selected_jd
        if jd_pasted_text:
            st.success(f"✅ JD loaded from library: **{selected_jd}** ({len(jd_pasted_text):,} chars)")

    # ── Readiness check ───────────────────────────────────────────────────────
    pasted_profiles  = st.session_state.get("pasted_profiles", [])
    skills_to_verify = st.session_state.get("skills_to_verify", [])
    has_profiles = bool(profile_files) or bool(pasted_profiles)
    has_jd       = bool(jd_file) or bool(jd_pasted_text.strip())

    missing = []
    if not api_key:      missing.append(f"{provider} API key (in AI Settings)")
    if not has_jd:       missing.append("Job Description")
    if not has_profiles: missing.append("at least one candidate profile")
    if missing:
        st.info("ℹ️ To begin, please provide: " + " · ".join(missing))

    rank_clicked = st.button("🚀  Evaluate & Rank Profiles",
                             disabled=bool(missing), use_container_width=True)

    # ── Run ranking ───────────────────────────────────────────────────────────
    if rank_clicked:
        with st.spinner("📄 Extracting Job Description…"):
            try:
                if jd_file:
                    jd_text = extract_text(jd_file)
                    eff_jd_name = jd_file.name.rsplit(".",1)[0]
                else:
                    jd_text = jd_pasted_text.strip()
                    eff_jd_name = jd_pasted_name
                if not jd_text.strip():
                    st.error("❌ Could not extract JD text."); return
            except Exception as e:
                st.error(f"❌ JD extraction failed: {e}"); return

        results       = []
        prog          = st.progress(0, text="Starting…")
        status_holder = st.empty()
        err_holder    = st.container()
        total         = len(profile_files) + len(pasted_profiles)
        counter       = 0
        def _run(text, name):
            return rank_profile(
                api_key, provider, model,
                jd_text, text, name,
                mandatory_reqs, additional_criteria,
                skills_to_verify
            )

        for pfile in profile_files:
            prog.progress(int(counter/total*100), text=f"Evaluating {counter+1}/{total}: {pfile.name}…")
            try:
                file_bytes = pfile.read(); pfile.seek(0)
                if pfile.name.lower().endswith(".pdf") and is_image_based_pdf(file_bytes):
                    err_holder.warning(
                        f"⚠️ **'{pfile.name}'** appears to be a **scanned/image-based PDF** — "
                        f"text cannot be extracted automatically. "
                        f"Please copy-paste the resume text using Step 3 (Paste Resume Text)."
                    )
                    results.append(_error_result(pfile.name, "Scanned PDF — no extractable text"))
                    counter += 1
                    continue
                pt = extract_text(pfile)
                if not pt.strip():
                    err_holder.warning(
                        f"⚠️ **'{pfile.name}'** — no text could be extracted. "
                        f"Try copy-pasting the resume using Step 3."
                    )
                    results.append(_error_result(pfile.name, "No text extracted"))
                    counter += 1
                    continue
                a = _run(pt, pfile.name)
                a["file"] = pfile.name
                results.append(a)
            except json.JSONDecodeError as e:
                err_holder.warning(f"⚠️ Skipped '{pfile.name}': unexpected AI format — {e}")
                results.append(_error_result(pfile.name, str(e)))
            except Exception as e:
                err_holder.warning(f"⚠️ Skipped '{pfile.name}': {e}")
                results.append(_error_result(pfile.name, str(e)))
            counter += 1

        for pp in pasted_profiles:
            dn = pp["name"]
            prog.progress(int(counter/total*100), text=f"Evaluating {counter+1}/{total}: {dn}…")
            try:
                if not pp["text"].strip(): raise ValueError("Empty pasted resume")
                a = _run(pp["text"], dn)
                a["file"] = dn
                results.append(a)
            except json.JSONDecodeError as e:
                err_holder.warning(f"⚠️ Skipped '{dn}': unexpected AI format — {e}")
                results.append(_error_result(dn, str(e)))
            except Exception as e:
                err_holder.warning(f"⚠️ Skipped '{dn}': {e}")
                results.append(_error_result(dn, str(e)))
            counter += 1

        prog.progress(100, text="✅ Done!")
        status_holder.success(f"{len(results)} profile(s) evaluated.")
        results.sort(key=lambda x: x.get("overall_score", 0), reverse=True)
        st.session_state["results"]  = results
        st.session_state["jd_name"]  = eff_jd_name

        # ── Append to screening history ───────────────────────────────────────
        for r in results:
            shared["history"].append({
                "timestamp":  datetime.now().strftime("%Y-%m-%d %H:%M"),
                "jd_name":    eff_jd_name,
                "candidate":  r.get("file","—").rsplit(".",1)[0] if "." in r.get("file","") else r.get("file","—"),
                "score":      r.get("overall_score", 0),
                "verdict":    r.get("verdict", "N/A"),
                "confidence": r.get("confidence_level", "N/A"),
                "seniority":  r.get("seniority_alignment", "N/A"),
                "must_have":  r.get("must_have_match_pct", 0),
                "experience": r.get("experience_match_pct", 0),
                "mandatory_met": r.get("mandatory_met", True),
                "strengths":  r.get("key_strengths", []),
                "gaps":       r.get("key_gaps", []),
                "explanation": r.get("final_explanation", ""),
            })

    # ── Display results ───────────────────────────────────────────────────────
    if not st.session_state.get("results"):
        return

    results = st.session_state["results"]
    jd_name = st.session_state.get("jd_name", "JD")

    st.divider()
    st.subheader("📊 Evaluation Results")

    total_r = len(results)
    avg_sc  = round(sum(r.get("overall_score",0) for r in results) / total_r)
    ss_cnt  = sum(1 for r in results if r.get("verdict")=="Strong Select")
    co_cnt  = sum(1 for r in results if r.get("verdict")=="Consider")
    rej_cnt = sum(1 for r in results if r.get("verdict")=="Reject")
    fm_cnt  = sum(1 for r in results if not r.get("mandatory_met",True))

    m1,m2,m3,m4,m5,m6 = st.columns(6)
    m1.metric("Profiles",          total_r)
    m2.metric("Avg Score",         f"{avg_sc}%")
    m3.metric("🟢 Strong Select",  ss_cnt)
    m4.metric("🟡 Consider",       co_cnt)
    m5.metric("🔴 Reject",         rej_cnt)
    m6.metric("⛔ Failed Mand.",   fm_cnt)

    st.divider()

    def _tick(v):
        s = str(v).lower()
        if s == "yes":     return "✅"
        if s == "no":      return "❌"
        if s == "partial": return "🟡"
        return "—"

    for i, r in enumerate(results):
        score      = r.get("overall_score", 0)
        verdict    = r.get("verdict", "N/A")
        confidence = r.get("confidence_level", "N/A")
        mand_ok    = r.get("mandatory_met", True)
        cp         = r.get("cv_profile", {})

        if not mand_ok:         ind = "🔴 Rejected — Mandatory Missing"
        elif verdict=="Strong Select": ind = "🟢 Strong Select"
        elif verdict=="Consider":      ind = "🟡 Consider"
        elif verdict=="Reject":        ind = "🔴 Reject"
        else:                          ind = "⚪ N/A"

        rank_emoji = ["🥇","🥈","🥉"][i] if i < 3 else f"#{i+1}"
        cname = r["file"].rsplit(".",1)[0] if "." in r["file"] else r["file"]

        with st.expander(
            f"{rank_emoji}  {cname}   |   Score: **{score}%**   |   {verdict}   "
            f"|   {confidence} confidence   |   {ind}",
            expanded=(i < 3)
        ):
            tab1, tab2, tab3 = st.tabs(
                ["👤 CV Profile", "🔍 Skills Checklist", "📊 Score & Assessment"]
            )

            # ── Tab 1: CV Profile ──────────────────────────────────────────
            with tab1:
                if cp:
                    p1, p2, p3 = st.columns(3)
                    p1.metric("Total Experience", cp.get("total_experience_label","—"))
                    p2.metric("Job Changes",       cp.get("job_changes_count","—"))
                    edu_list = cp.get("education",[])
                    p3.metric("Qualifications",    len(edu_list))

                    st.markdown("---")
                    c_left, c_right = st.columns(2)

                    with c_left:
                        st.markdown("##### 💼 Work History")
                        roles = cp.get("roles", [])
                        if roles:
                            for role in roles:
                                short = role.get("short_tenure", False)
                                flag  = " ⚠️" if short else ""
                                st.markdown(
                                    f"**{role.get('title','—')}** @ {role.get('company','—')}{flag}  \n"
                                    f"_{role.get('start','?')} – {role.get('end','?')}  "
                                    f"({role.get('duration_months','?')} months)_"
                                )
                        else:
                            st.caption("No work history extracted.")

                        st.markdown("##### 🎓 Education")
                        if edu_list:
                            for e in edu_list:
                                st.markdown(
                                    f"**{e.get('degree','')} {e.get('field','')}**  \n"
                                    f"{e.get('institution','—')} · {e.get('year','—')}"
                                )
                        else:
                            st.caption("No education extracted.")

                    with c_right:
                        st.markdown("##### ⏸️ Career Gaps")
                        gaps = cp.get("career_gaps", [])
                        if gaps:
                            for g in gaps:
                                note = f" — {g['note']}" if g.get("note") else ""
                                st.warning(
                                    f"**{g.get('from','?')} → {g.get('to','?')}**  "
                                    f"({g.get('duration_months','?')} months){note}"
                                )
                        else:
                            st.success("No significant career gaps identified.")

                        st.markdown("##### ⚠️ Risk Flags")
                        rf_list = r.get("risk_flags", [])
                        if rf_list and rf_list != ["None identified"]:
                            for rf in rf_list:
                                st.error(f"• {rf}")
                        else:
                            st.success("No risk flags.")
                else:
                    st.info("CV profile data not available for this candidate.")

            # ── Tab 2: Skills Checklist ────────────────────────────────────
            with tab2:
                sc_list = r.get("skills_checklist", [])
                if sc_list:
                    # Deduplicate: if same skill appears twice, keep the better status
                    status_rank = {"Present": 0, "Partial": 1, "Absent": 2}
                    seen = {}
                    for s in sc_list:
                        skill_key = s.get("skill","").lower().strip()
                        if skill_key not in seen:
                            seen[skill_key] = s
                        else:
                            existing = seen[skill_key]
                            if status_rank.get(s.get("status","Absent"), 2) < status_rank.get(existing.get("status","Absent"), 2):
                                seen[skill_key] = s
                    sc_list_deduped = list(seen.values())

                    st.markdown("##### 🔍 Skills Verification Results")
                    rows = []
                    for s in sc_list_deduped:
                        status = s.get("status","—")
                        if status == "Present":
                            icon = "✅ Present"
                        elif status == "Partial":
                            icon = "🟡 Partial"
                        else:
                            icon = "❌ Absent"
                        rows.append({
                            "Skill":     s.get("skill","—"),
                            "Status":    icon,
                            "Evidence":  s.get("evidence","—"),
                        })

                    st.markdown("| Skill | Status | Evidence |")
                    st.markdown("|:---|:---:|:---|")
                    for row in rows:
                        st.markdown(f"| {row['Skill']} | {row['Status']} | {row['Evidence']} |")

                    present  = sum(1 for s in sc_list_deduped if s.get("status")=="Present")
                    partial  = sum(1 for s in sc_list_deduped if s.get("status")=="Partial")
                    absent   = sum(1 for s in sc_list_deduped if s.get("status")=="Absent")
                    st.markdown(f"\n✅ **{present} Present** &nbsp;|&nbsp; 🟡 **{partial} Partial** &nbsp;|&nbsp; ❌ **{absent} Absent**")
                else:
                    if skills_to_verify:
                        st.info("Skills checklist was not returned by the AI for this candidate.")
                    else:
                        st.info("No skills to verify were specified. Add them in the sidebar → **Skills to Verify**.")

                # Key strengths and gaps here too
                st.markdown("---")
                sg1, sg2 = st.columns(2)
                with sg1:
                    st.markdown("##### ✅ Key Strengths")
                    for s in r.get("key_strengths", []):
                        st.markdown(f"• {s}")
                with sg2:
                    st.markdown("##### ❌ Key Gaps")
                    for g in r.get("key_gaps", []):
                        st.markdown(f"• {g}")

            # ── Tab 3: Score & Assessment ──────────────────────────────────
            with tab3:
                sc1, sc2 = st.columns([1, 2])
                with sc1:
                    st.markdown("##### 📊 Score Breakdown")
                    st.progress(score / 100)
                    st.markdown(f"**Overall: {score} / 100**")

                    mh  = r.get("must_have_match_pct","—")
                    gth = r.get("good_to_have_match_pct","—")
                    exp = r.get("experience_match_pct","—")
                    add = r.get("additional_criteria_match_pct","—")

                    st.markdown(f"""
| Criterion | Weight | Score |
|:---|:---:|---:|
| Must-Have Skills (JD) | 35% | {mh}% |
| Experience Match | 25% | {exp}% |
| Additional Criteria | 25% | {add}% |
| Good-to-Have Skills (JD) | 15% | {gth}% |
""")
                    st.markdown(f"**Verdict:** `{verdict}`   **Confidence:** `{confidence}`")
                    st.markdown(f"**Seniority:** `{r.get('seniority_alignment','N/A')}`")
                    st.markdown(
                        f"Role: `{r.get('role_relevance','N/A')}` · "
                        f"Domain: `{r.get('domain_relevance','N/A')}` · "
                        f"Years: `{r.get('years_match','N/A')}`"
                    )

                    if not mand_ok:
                        st.error("⛔ Missing Mandatory Requirements:")
                        for m in r.get("mandatory_missing", []):
                            st.markdown(f"  ❌ {m}")

                with sc2:
                    penalties = r.get("penalties_applied", [])
                    if penalties:
                        st.markdown("##### 🔻 Penalties Applied")
                        for p in penalties:
                            st.markdown(f"• {p}")

                    add_match = r.get("additional_criteria_match", [])
                    if add_match and add_match != ["No additional criteria provided"]:
                        st.markdown("##### 🎯 Additional Criteria")
                        for ac in add_match:
                            st.markdown(f"• {ac}")

                st.markdown("##### 💬 Final Explanation")
                st.info(r.get("final_explanation","—"))


    # ── Download current results ───────────────────────────────────────────────
    st.divider()
    excel_bytes = generate_excel(results, jd_name)
    st.download_button(
        label="📥 Download Excel Report",
        data=excel_bytes,
        file_name=f"Profile_Rankings_{jd_name}_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

    # ══════════════════════════════════════════════════════════════════════════
    # SCREENING HISTORY
    # ══════════════════════════════════════════════════════════════════════════
    history = shared["history"]
    if history:
        st.divider()
        st.subheader("📚 Screening History")
        st.caption(f"{len(history)} candidate(s) screened across all JDs")

        # ── Filters ───────────────────────────────────────────────────────────
        hf1, hf2, hf3 = st.columns(3)
        with hf1:
            jd_options = ["All JDs"] + sorted(set(h["jd_name"] for h in history))
            filter_jd = st.selectbox("Filter by JD", jd_options, key="hist_filter_jd")
        with hf2:
            verdict_options = ["All Verdicts", "Strong Select", "Consider", "Reject"]
            filter_verdict = st.selectbox("Filter by Verdict", verdict_options, key="hist_filter_verdict")
        with hf3:
            search_name = st.text_input("Search candidate", placeholder="Type name…", key="hist_search")

        filtered = history
        if filter_jd != "All JDs":
            filtered = [h for h in filtered if h["jd_name"] == filter_jd]
        if filter_verdict != "All Verdicts":
            filtered = [h for h in filtered if h["verdict"] == filter_verdict]
        if search_name.strip():
            filtered = [h for h in filtered if search_name.lower() in h["candidate"].lower()]

        filtered_sorted = sorted(filtered, key=lambda x: x["score"], reverse=True)

        # ── Summary table ─────────────────────────────────────────────────────
        st.markdown(f"**Showing {len(filtered_sorted)} of {len(history)} records**")

        # Table header
        st.markdown("""
| # | Candidate | JD | Score | Verdict | Must-Have % | Experience % | Screened On |
|:--|:----------|:---|------:|:--------|------------:|-------------:|:------------|""")

        for idx, h in enumerate(filtered_sorted, 1):
            v = h["verdict"]
            if v == "Strong Select":   vbadge = "🟢 Strong Select"
            elif v == "Consider":      vbadge = "🟡 Consider"
            elif v == "Reject":        vbadge = "🔴 Reject"
            else:                      vbadge = "⚪ N/A"
            if not h.get("mandatory_met", True): vbadge = "⛔ Rejected"

            st.markdown(
                f"| {idx} | **{h['candidate']}** | {h['jd_name']} | "
                f"**{h['score']}%** | {vbadge} | {h['must_have']}% | "
                f"{h['experience']}% | {h['timestamp']} |"
            )

        # ── Expandable detail per candidate ───────────────────────────────────
        st.markdown("---")
        st.markdown("##### Details")
        for h in filtered_sorted:
            v = h["verdict"]
            badge = "🟢" if v=="Strong Select" else ("🟡" if v=="Consider" else "🔴")
            with st.expander(f"{badge} {h['candidate']}  —  {h['score']}%  |  {h['jd_name']}  |  {h['timestamp']}"):
                d1, d2 = st.columns(2)
                with d1:
                    st.markdown(f"**Verdict:** {h['verdict']}  \n**Confidence:** {h['confidence']}  \n**Seniority:** {h['seniority']}")
                    if h.get("strengths"):
                        st.markdown("**Strengths:**")
                        for s in h["strengths"]: st.markdown(f"  • {s}")
                with d2:
                    if h.get("gaps"):
                        st.markdown("**Gaps:**")
                        for g in h["gaps"]: st.markdown(f"  • {g}")
                if h.get("explanation"):
                    st.info(h["explanation"])

        # ── Export / Import history ────────────────────────────────────────────
        st.markdown("---")
        exp_col, imp_col, clr_col = st.columns(3)

        with exp_col:
            history_json = json.dumps(shared["history"], indent=2)
            st.download_button(
                "📥 Export History (JSON)",
                data=history_json,
                file_name=f"screening_history_{datetime.now().strftime('%Y%m%d')}.json",
                mime="application/json",
                use_container_width=True,
            )
        with imp_col:
            uploaded_hist = st.file_uploader("📤 Import History (JSON)",
                                              type=["json"], key="import_history_uploader")
            if uploaded_hist:
                try:
                    imported = json.loads(uploaded_hist.read())
                    existing_keys = {(h["candidate"], h["timestamp"]) for h in shared["history"]}
                    added = 0
                    for entry in imported:
                        if (entry.get("candidate"), entry.get("timestamp")) not in existing_keys:
                            shared["history"].append(entry)
                            added += 1
                    st.success(f"✅ Imported {added} new record(s)")
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Failed to import: {e}")
        with clr_col:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🗑️ Clear History", key="clear_history_btn", use_container_width=True):
                shared["history"].clear()
                st.rerun()

        # ── JD Library export/import ───────────────────────────────────────────
        if shared["jd_library"]:
            st.markdown("---")
            st.markdown("##### 📂 JD Library")
            st.caption(f"{len(shared['jd_library'])} JD(s) saved")
            jl_exp, jl_imp = st.columns(2)
            with jl_exp:
                jd_lib_json = json.dumps(shared["jd_library"], indent=2)
                st.download_button(
                    "📥 Export JD Library (JSON)",
                    data=jd_lib_json,
                    file_name=f"jd_library_{datetime.now().strftime('%Y%m%d')}.json",
                    mime="application/json",
                    use_container_width=True,
                )
            with jl_imp:
                uploaded_jdlib = st.file_uploader("📤 Import JD Library (JSON)",
                                                   type=["json"], key="import_jdlib_uploader")
                if uploaded_jdlib:
                    try:
                        imported_lib = json.loads(uploaded_jdlib.read())
                        shared["jd_library"].update(imported_lib)
                        st.success(f"✅ Loaded {len(imported_lib)} JD(s)")
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ Failed to import JD library: {e}")


if __name__ == "__main__":
    main()
